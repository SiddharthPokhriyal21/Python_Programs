import os
import PyPDF2
import docx2pdf
import pyttsx3
import fitz
from pptx import Presentation
from pyexcelerate import Workbook
from docx import Document
from docx.shared import Inches


# def extract_text_and_images(pdf_file):
#     doc = fitz.open(pdf_file)
#     text = ""
#     images = []

#     for page_num in range(len(doc)):
#         page = doc.load_page(page_num)
        
#         # Extract text
#         text += page.get_text()

#         # Extract images
#         image_list = page.get_images(full=True)
#         for image_index, img in enumerate(image_list):
#             base_image = doc.extract_image(image_index)
#             image_bytes = base_image["image"]
#             image_extension = base_image["ext"]
#             images.append((image_bytes, image_extension))
    
#     return text, images

# def add_text_to_doc(text, doc):
#     paragraphs = text.split('\n')
#     for paragraph in paragraphs:
#         doc.add_paragraph(paragraph)

# def add_images_to_doc(images, doc):
#     for image_bytes, image_extension in images:
#         image_file = f'image.{image_extension}'
#         with open(image_file, 'wb') as f:
#             f.write(image_bytes)
#         doc.add_picture(image_file, width=Inches(4))
#         # Delete temporary image file
#         os.remove(image_file)


if __name__ == '__main__':
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    engine.setProperty('voice', voices[1].id)
    
    print("Welcome to PDF Modifier 1.0")
    engine.say("Welcome to PDF Modifier 1.0")
    engine.runAndWait()

    while True:
        print("What do you want to do?")
        engine.say("What do you want to do?")
        engine.runAndWait()
        x = int(input("Enter 1 to merge PDFs.\nEnter 2 to convert a Word file to a PDF.\nEnter 3 to convert a PDF to a Word file. (UNDER PROGRESS)\nEnter 4 to convert an PowerPoint presentation to a PDF. (UNDER PROGRESS)\nEnter 5 to convert an Excel sheet to a PDF. (UNDER PROGRESS)\nEnter 0 to exit.\n"))

        if x == 0:
            print('Thank you!')
            engine.say('Thank you!')
            engine.runAndWait()
            break

        elif x == 1:
            pdf_files = []
            while True:
                path = input("Enter the source path: ")
                if path == '?':
                    break
                pdf_files.append(path)

            merger = PyPDF2.PdfMerger()

            for file_name in pdf_files:
                pdf_file = open(file_name, 'rb')
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                merger.append(pdf_reader)
            pdf_file.close()

            dest = input("Enter the destination path: ")
            merger.write(dest)
        
        elif x == 2:
            word_files = []
            while True:
                path = input("Enter the source path: ")
                if path == '?':
                    break
                word_files.append(path)

            dest = input("Enter the destination path: ")
            for file in word_files:
                docx2pdf.convert(file, dest)
        
        # elif x == 3:
        #     pdf_files = []
        #     while True:
        #         path = input("Enter the source path: ")
        #         if path == '?':
        #             break
        #         pdf_files.append(path)

        #     dest = input("Enter the destination path: ")
        #     for file in pdf_files:
        #         text, images = extract_text_and_images(file)
        #         doc = Document()
        #         add_text_to_doc(text, doc)
        #         add_images_to_doc(images, doc)
        #         doc.save(dest)

        # elif x == 4:
        #     powerpoint_files = []
        #     while True:
        #         path = input("Enter the source path: ")
        #         if path == '?':
        #             break
        #         powerpoint_files.append(path)

        #     dest = input("Enter the destination path: ")
        #     for file in powerpoint_files:
        #         ppt = Presentation(file)
        #         ppt.save(dest)

        # elif x == 5:
        #     excel_files = []
        #     while True:
        #         path = input("Enter the source path: ")
        #         if path == '?':
        #             break
        #         excel_files.append(path)

        #     dest = input("Enter the destination path: ")
        #     for file in excel_files:
        #         wb = Workbook()
        #         wb.read(file)
        #         wb.save(dest)

        else:
            print("Kindly try a valid input.")
            engine.say("Kindly try a valid input.")
            engine.runAndWait()